import pandas as pd

# Load R export
cox_table = pd.read_csv("cox_summary.csv")

# Preview
print(cox_table.head())

from docx import Document
from docx.shared import Inches

# Create a new Word doc
doc = Document()
doc.add_heading("Healthcare Survival Analysis Report", 0)

# Add Cox table
doc.add_heading("Cox Proportional Hazards Model", level=1)
doc.add_paragraph("Summary table of hazard ratios and significance.")
doc.add_table(rows=1, cols=len(cox_table.columns))

# Fill table
hdr_cells = doc.tables[0].rows[0].cells
for i, col in enumerate(cox_table.columns):
    hdr_cells[i].text = col

for _, row in cox_table.iterrows():
    cells = doc.tables[0].add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(val) 

# Add survival plot
doc.add_heading("Kaplan-Meier Curve", level=1)
doc.add_picture("survival_curve.png", width=Inches(5))

# Save
doc.save("Healthcare_Survival_Report.docx")

